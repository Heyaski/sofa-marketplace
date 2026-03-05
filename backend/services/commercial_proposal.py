"""
Сервис генерации коммерческого предложения (КП) в формате PDF и DOCX.
Формирует таблицу с фото, ценой, габаритами товаров из корзины.
"""
import io
import os
import re
import tempfile
import requests
from datetime import datetime
from urllib.parse import quote_plus

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
    Image as RLImage, PageBreak, Frame, PageTemplate
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage

from django.conf import settings

# Имя шрифта, зарегистрированного для КП
_CP_FONT = 'CPFont'
_CP_FONT_BOLD = 'CPFont-Bold'


def _strip_brand(title, brand):
    """
    Убирает бренд из названия товара и оставляет только
    тип мебели и цвет, если удаётся распарсить.
    Пример:
      "Табурет мягкий Handy светло-коричневого цвета" + brand "Handy"
      → "Табурет светло-коричневого цвета"
    """
    base = title or ''
    if brand and brand.strip():
        escaped = re.escape(brand.strip())
        pattern = re.compile(r'\s*' + escaped + r'\s*', re.IGNORECASE)
        base = pattern.sub(' ', base)
    base = re.sub(r'\s+', ' ', base).strip()
    if not base:
        return ''

    # Первый "тип" мебели — первое слово до пробела/запятой
    m_type = re.match(r'^\s*([^\s,]+)', base)
    item_type = m_type.group(1) if m_type else ''

    # Хвост с цветом — одно/составное слово непосредственно перед "цвет..." в конце строки
    m_color = re.search(r'((?:[А-Яа-яЁё]+-)*[А-Яа-яЁё]+\s+цвет[а-я]*)\s*$', base, re.IGNORECASE)
    if item_type and m_color:
        color_part = m_color.group(1).strip()
        return f"{item_type} {color_part}".strip()

    return base



def register_fonts():
    """Регистрирует инженерный шрифт (ISOCPEUR → Courier New → built-in Courier)."""
    font_candidates = [
        # ISOCPEUR — инженерный шрифт (Windows, AutoCAD)
        'C:/Windows/Fonts/isocpeur.ttf',
        # GOST type B (если установлен)
        'C:/Windows/Fonts/GOST_B.ttf',
        # Courier New (Windows)
        'C:/Windows/Fonts/cour.ttf',
        # Courier New (Linux — пакет ttf-mscorefonts-installer)
        '/usr/share/fonts/truetype/msttcorefonts/Courier_New.ttf',
        '/usr/share/fonts/truetype/courier-prime/CourierPrime-Regular.ttf',
        # DejaVu Mono (Linux fallback)
        '/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf',
        '/Library/Fonts/Courier New.ttf',
    ]
    bold_candidates = [
        'C:/Windows/Fonts/courbd.ttf',
        '/usr/share/fonts/truetype/msttcorefonts/Courier_New_Bold.ttf',
        '/usr/share/fonts/truetype/courier-prime/CourierPrime-Bold.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf',
        '/Library/Fonts/Courier New Bold.ttf',
    ]

    registered_regular = False
    for path in font_candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(_CP_FONT, path))
                registered_regular = True
                break
            except Exception:
                pass

    for path in bold_candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(_CP_FONT_BOLD, path))
                break
            except Exception:
                pass

    return registered_regular


def get_font_name(bold=False):
    """Возвращает имя зарегистрированного шрифта; фолбэк — встроенный Courier."""
    target = _CP_FONT_BOLD if bold else _CP_FONT
    try:
        pdfmetrics.getFont(target)
        return target
    except KeyError:
        # Courier — встроенный инженерный шрифт PDF, всегда доступен
        return 'Courier-Bold' if bold else 'Courier'


def _read_file_field(field_file):
    """
    Читает содержимое файлового поля Django (ImageField/FileField).
    Работает и с локальным хранилищем, и с S3.
    Возвращает io.BytesIO или None.
    """
    if not field_file:
        return None
    
    # Способ 1: Попробовать прочитать через Django storage API (.open / .read)
    try:
        field_file.open('rb')
        data = field_file.read()
        field_file.close()
        if data:
            return io.BytesIO(data)
    except Exception:
        pass
    
    # Способ 2: Если есть URL — скачать по HTTP (S3, CDN и т.д.)
    try:
        url = field_file.url
        if url:
            # Для относительных URL добавляем домен (на случай локального dev-сервера)
            if url.startswith('/'):
                url = f"http://localhost:8000{url}"
            resp = requests.get(url, timeout=15)
            if resp.status_code == 200 and resp.content:
                return io.BytesIO(resp.content)
    except Exception:
        pass
    
    # Способ 3: Если есть локальный путь — прочитать с диска
    try:
        path = field_file.path
        if os.path.exists(path):
            with open(path, 'rb') as f:
                return io.BytesIO(f.read())
    except Exception:
        pass
    
    return None


def get_product_image(product):
    """Получает изображение товара для КП как io.BytesIO.
    Использует уже загруженные фото товара (ProductImage, FileAsset, image, photo_url).
    Работает с любым storage backend (локальный, S3 и т.д.).
    """
    # Приоритет 1: Первое изображение из ProductImage (загружены через импорт/админку)
    if product.images.exists():
        first_image = product.images.first()
        if first_image and first_image.image:
            result = _read_file_field(first_image.image)
            if result:
                return result
    
    # Приоритет 2: Изображения из FileAsset по ID
    image_assets = product.get_image_assets()
    if image_assets.exists():
        first_asset = image_assets.first()
        if first_asset and first_asset.file:
            result = _read_file_field(first_asset.file)
            if result:
                return result
    
    # Приоритет 3: photo_url (из Excel импорта — внешняя ссылка)
    if product.photo_url:
        try:
            resp = requests.get(product.photo_url, timeout=15)
            if resp.status_code == 200 and resp.content:
                return io.BytesIO(resp.content)
        except Exception:
            pass
    
    # Приоритет 4: Основное изображение товара (старое поле image)
    if product.image:
        result = _read_file_field(product.image)
        if result:
            return result
    
    return None


def format_dimensions(product):
    """Форматирует габариты товара для КП"""
    dims = []
    if product.width:
        dims.append(str(int(product.width)))
    if product.depth:
        dims.append(str(int(product.depth)))
    if product.height:
        dims.append(str(int(product.height)))
    
    if dims:
        return f"Габариты (ш*г*в):\n{'x'.join(dims)}"
    return ""


def generate_commercial_proposal_pdf(proposal_request):
    """
    Генерирует PDF коммерческого предложения.
    
    Args:
        proposal_request: CommercialProposalRequest instance
    
    Returns:
        bytes: PDF файл в виде байтов
    """
    register_fonts()
    
    font_name = get_font_name()
    font_name_bold = get_font_name(bold=True)
    
    buffer = io.BytesIO()
    
    # Создаем документ в альбомной ориентации A4
    page_width, page_height = landscape(A4)
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=25 * mm,
    )
    
    # Стили
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CPTitle',
        parent=styles['Title'],
        fontName=font_name_bold,
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=10,
        leading=20,
    )
    
    header_style = ParagraphStyle(
        'CPHeader',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        alignment=TA_LEFT,
        spaceAfter=3,
        leading=14,
    )
    
    header_bold_style = ParagraphStyle(
        'CPHeaderBold',
        parent=styles['Normal'],
        fontName=font_name_bold,
        fontSize=10,
        alignment=TA_LEFT,
        spaceAfter=3,
        leading=14,
    )
    
    cell_style = ParagraphStyle(
        'CPCell',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=8,
        alignment=TA_CENTER,
        leading=10,
    )
    
    cell_left_style = ParagraphStyle(
        'CPCellLeft',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=8,
        alignment=TA_LEFT,
        leading=10,
    )
    
    cell_small_style = ParagraphStyle(
        'CPCellSmall',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=7,
        alignment=TA_CENTER,
        leading=9,
        wordWrap='CJK',
    )
    
    header_cell_style = ParagraphStyle(
        'CPHeaderCell',
        parent=styles['Normal'],
        fontName=font_name_bold,
        fontSize=8,
        alignment=TA_CENTER,
        leading=10,
    )
    
    note_style = ParagraphStyle(
        'CPNote',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=7,
        alignment=TA_LEFT,
        leading=9,
    )
    
    # Элементы для документа
    elements = []
    
    basket = proposal_request.basket
    items = basket.items.select_related('product').all()
    
    # === ЗАГОЛОВОК ===
    elements.append(Paragraph("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", title_style))
    elements.append(Spacer(1, 5 * mm))
    
    # === МЕТА-ИНФОРМАЦИЯ ===
    # Форматируем дату
    date_str = proposal_request.created_at.strftime('%d %B %Y г.')
    # Русские месяцы
    months_ru = {
        'January': 'января', 'February': 'февраля', 'March': 'марта',
        'April': 'апреля', 'May': 'мая', 'June': 'июня',
        'July': 'июля', 'August': 'августа', 'September': 'сентября',
        'October': 'октября', 'November': 'ноября', 'December': 'декабря',
    }
    date_str = proposal_request.created_at.strftime('%d {month} %Y г.')
    month_en = proposal_request.created_at.strftime('%B')
    date_str = date_str.replace('{month}', months_ru.get(month_en, month_en))
    
    elements.append(Paragraph(f"<b>Дата:</b> {date_str}", header_style))
    elements.append(Paragraph(f"<b>Клиент:</b> {proposal_request.client_name}", header_style))
    
    if proposal_request.company_name:
        elements.append(Paragraph(f"<b>От:</b> {proposal_request.company_name}", header_style))
    
    elements.append(Paragraph(f"<b>Проект:</b> {proposal_request.project_name}", header_style))
    elements.append(Spacer(1, 8 * mm))
    
    # === ТАБЛИЦА ТОВАРОВ ===
    # Заголовки таблицы
    table_header = [
        Paragraph('ID', header_cell_style),
        Paragraph('Наименование', header_cell_style),
        Paragraph('Изображение', header_cell_style),
        Paragraph('Кол-во,&nbsp;шт', header_cell_style),
        Paragraph('Цена за<br/>шт., руб.', header_cell_style),
        Paragraph('Сумма,<br/>руб.', header_cell_style),
        Paragraph('Магазин,<br/>ссылка', header_cell_style),
        Paragraph('Примечание', header_cell_style),
    ]
    
    table_data = [table_header]
    
    total_sum = 0
    
    for idx, item in enumerate(items, start=1):
        product = item.product
        quantity = item.quantity
        price = float(product.price)
        item_total = price * quantity
        total_sum += item_total
        
        # ID — берём ID 3D модели из поля model_3d_asset_ids
        model_id = product.model_3d_asset_ids.strip().split(',')[0] if product.model_3d_asset_ids and product.model_3d_asset_ids.strip() else f'#{product.id}'
        item_id = Paragraph(model_id, cell_small_style)
        
        # Наименование (без бренда)
        display_title = _strip_brand(product.title, product.brand)
        item_name = Paragraph(display_title, cell_left_style)
        
        # Изображение (сохраняем пропорции, вписываем в ячейку)
        img_data = get_product_image(product)
        if img_data:
            try:
                max_w, max_h = 65, 65  # points; чуть меньше ширины колонки (75pt) и высоты строки (75pt)
                img_data.seek(0)
                reader = ImageReader(img_data)
                iw, ih = reader.getSize()
                if iw > 0 and ih > 0:
                    ratio = min(max_w / iw, max_h / ih)
                    img_data.seek(0)
                    img = RLImage(img_data, width=iw * ratio, height=ih * ratio)
                    img.hAlign = 'CENTER'
                    item_image = img
                else:
                    item_image = Paragraph('—', cell_style)
            except Exception:
                item_image = Paragraph('—', cell_style)
        else:
            item_image = Paragraph('—', cell_style)
        
        # Количество
        item_qty = Paragraph(str(quantity), cell_style)
        
        # Цена
        item_price = Paragraph(f'{price:,.0f}'.replace(',', ' '), cell_style)
        
        # Сумма
        item_sum = Paragraph(f'{item_total:,.0f}'.replace(',', ' '), cell_style)
        
        # Ссылка — поиск по оригинальному названию (с брендом) для точных результатов
        original_title = product.title or display_title
        search_query = quote_plus(f'{original_title} купить')
        search_url = f'https://ya.ru/search/?text={search_query}'
        short_label = f'ya.ru: {display_title[:30]}...' if len(display_title) > 30 else f'ya.ru: {display_title}'
        item_shop = Paragraph(
            f'<a href="{search_url}" color="blue">{short_label}</a>',
            cell_small_style,
        )
        
        # Примечание (габариты + доп. информация)
        notes_parts = []
        dims = format_dimensions(product)
        if dims:
            notes_parts.append(dims)
        
        if product.cp_notes:
            notes_parts.append(product.cp_notes)
        elif product.brand:
            notes_parts.append(f"Производитель: {product.brand}")
        
        item_notes = Paragraph('<br/>'.join(notes_parts) if notes_parts else '—', cell_small_style)
        
        table_data.append([
            item_id, item_name, item_image, item_qty, 
            item_price, item_sum, item_shop, item_notes
        ])
    
    # Определяем ширины столбцов
    available_width = page_width - 30 * mm  # margins
    col_widths = [
        60,   # ID (model_3d_asset_ids, например «Пуф1497»)
        85,   # Наименование
        75,   # Изображение
        45,   # Кол-во, шт
        55,   # Цена
        55,   # Сумма
        80,   # Магазин
        available_width - 60 - 85 - 75 - 45 - 55 - 55 - 80,  # Примечание
    ]
    
    # Минимальная высота строк
    row_heights = [25]  # заголовок
    for _ in items:
        row_heights.append(75)  # строки с товарами
    
    table = Table(table_data, colWidths=col_widths, rowHeights=row_heights)
    
    table.setStyle(TableStyle([
        # Заголовок
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.85, 0.85, 0.85)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('FONTNAME', (0, 0), (-1, 0), font_name_bold),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        
        # Все ячейки
        ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.6, 0.6, 0.6)),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # ID - по центру
        ('ALIGN', (2, 1), (2, -1), 'CENTER'),  # Изображение - по центру
        ('ALIGN', (3, 1), (5, -1), 'CENTER'),  # Кол-во, Цена, Сумма - по центру
        
        # Чередование фона строк
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.Color(0.97, 0.97, 0.97)]),
        
        # Padding
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 5 * mm))
    
    # === ИТОГО ===
    total_style = ParagraphStyle(
        'CPTotal',
        parent=styles['Normal'],
        fontName=font_name_bold,
        fontSize=12,
        alignment=TA_RIGHT,
        spaceAfter=5,
    )
    elements.append(Paragraph(
        f"ИТОГО: {total_sum:,.0f} руб.".replace(',', ' '),
        total_style
    ))
    elements.append(Spacer(1, 8 * mm))
    
    # === ПРИМЕЧАНИЯ ===
    elements.append(Paragraph("<b>Примечания:</b>", header_bold_style))
    elements.append(Paragraph(
        "1. Смотреть совместно с планом расстановки мебели и развертками.",
        note_style
    ))
    elements.append(Paragraph(
        "2. Детальные чертежи для мебели индивидуального производства составлять совместно с поставщиками.",
        note_style
    ))
    elements.append(Spacer(1, 10 * mm))
    
    # === КАРТОЧКА ПРОЕКТА (нижний правый угол) ===
    card_data = [
        [Paragraph('<b>КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ</b>', ParagraphStyle(
            'CardTitle', fontName=font_name_bold, fontSize=8, alignment=TA_CENTER, leading=10
        ))],
        [Paragraph(f'Дизайн-проект: {proposal_request.project_name}', ParagraphStyle(
            'CardProject', fontName=font_name, fontSize=7, alignment=TA_CENTER, leading=9
        ))],
    ]
    
    card_table = Table(card_data, colWidths=[150])
    card_table.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 1, colors.black),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.6, 0.6, 0.6)),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    
    # Обертка для позиционирования карточки справа
    wrapper_data = [['', card_table]]
    wrapper_table = Table(wrapper_data, colWidths=[available_width - 160, 160])
    wrapper_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    
    elements.append(wrapper_table)
    
    # Собираем PDF
    doc.build(elements)
    
    buffer.seek(0)
    return buffer.getvalue()


def _set_run_font(run, font_name, size_pt=None, bold=None):
    """Применяет инженерный шрифт к run."""
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _add_hyperlink_to_cell(cell, url, text, font_name, font_size=8):
    """Вставляет гиперссылку в ячейку таблицы DOCX."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = cell.paragraphs[0]
    paragraph.clear()
    # Принудительно задаём выравнивание по центру через XML
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:jc')):
        pPr.remove(old)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True,
        )

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')

        # Шрифт
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rpr.append(rFonts)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))  # half-points
        rpr.append(sz)

        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rpr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rpr.append(u)

        new_run.append(rpr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        paragraph.add_run(text)


def _set_cell_font(cell, font_name, font_size=9):
    """Применяет шрифт ко всем run-ам в ячейке."""
    for para in cell.paragraphs:
        for run in para.runs:
            _set_run_font(run, font_name, font_size)


def generate_commercial_proposal_docx(proposal_request):
    """
    Генерирует DOCX коммерческого предложения — 1 в 1 как PDF.
    Альбомная A4, серый заголовок, чередование строк, картинки, карточка проекта.
    """
    from docx.shared import Inches
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement

    basket = proposal_request.basket
    items = basket.items.select_related('product').all()

    doc = Document()

    # === СТРАНИЦА: альбомная A4, поля как в PDF (15мм / 25мм низ) ===
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2.5)

    # Courier New — тот же инженерный шрифт, что и в PDF (доступен на Windows/Linux)
    docx_font = "Courier New"

    normal_style = doc.styles["Normal"]
    normal_style.font.name = docx_font
    normal_style.font.size = Pt(9)

    # ---- хелперы ----

    def _par(text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        _set_run_font(p.add_run(text), docx_font, size, bold)
        return p

    def _par_labeled(label, value, size=10):
        p = doc.add_paragraph()
        _set_run_font(p.add_run(f"{label}: "), docx_font, size, bold=True)
        _set_run_font(p.add_run(value), docx_font, size, bold=False)
        return p

    def _cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
        para = cell.paragraphs[0]
        para.clear()
        # Принудительно задаём выравнивание через XML (paragraph.alignment иногда не срабатывает в ячейках)
        align_val = 'center' if align == WD_ALIGN_PARAGRAPH.CENTER else 'left'
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(_qn('w:jc')):
            pPr.remove(old)
        jc = _OxmlElement('w:jc')
        jc.set(_qn('w:val'), align_val)
        pPr.append(jc)
        _set_run_font(para.add_run(text), docx_font, size, bold)

    def _cell_bg(cell, hex_color):
        """Серый/белый фон ячейки."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # удалить старый shd если есть
        for old in tcPr.findall(_qn('w:shd')):
            tcPr.remove(old)
        shd = _OxmlElement('w:shd')
        shd.set(_qn('w:val'), 'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _remove_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = _OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = _OxmlElement(f'w:{side}')
            el.set(_qn('w:val'), 'none')
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _table_full_width(tbl_obj):
        """Растянуть таблицу на 100% ширины страницы."""
        tblPr = tbl_obj._tbl.tblPr
        for old in tblPr.findall(_qn('w:tblW')):
            tblPr.remove(old)
        tblW = _OxmlElement('w:tblW')
        tblW.set(_qn('w:w'), '5000')
        tblW.set(_qn('w:type'), 'pct')
        tblPr.append(tblW)

    # ---- ЗАГОЛОВОК ----
    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title_par.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ"), docx_font, 16, bold=True)
    doc.add_paragraph()

    # ---- МЕТА (с жирными метками, как в PDF) ----
    months_ru = {
        'January': 'января', 'February': 'февраля', 'March': 'марта',
        'April': 'апреля', 'May': 'мая', 'June': 'июня',
        'July': 'июля', 'August': 'августа', 'September': 'сентября',
        'October': 'октября', 'November': 'ноября', 'December': 'декабря',
    }
    month_en = proposal_request.created_at.strftime('%B')
    date_str = proposal_request.created_at.strftime('%d {month} %Y г.').replace(
        '{month}', months_ru.get(month_en, month_en)
    )
    _par_labeled("Дата", date_str)
    _par_labeled("Клиент", proposal_request.client_name)
    if proposal_request.company_name:
        _par_labeled("От", proposal_request.company_name)
    _par_labeled("Проект", proposal_request.project_name)
    doc.add_paragraph()

    # ---- ТАБЛИЦА ТОВАРОВ ----
    # Ширины в Cm пропорциональны PDF (доступно ~26.7 см)
    # PDF: 60 / 85 / 75 / 45 / 55 / 55 / 80 / ~305 pts
    col_cms = [2.12, 3.00, 2.65, 1.59, 1.94, 1.94, 2.83, 10.63]

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    _table_full_width(table)

    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = Cm(col_cms[idx])

    # Заголовок таблицы — серый фон D9D9D9 (= 0.85 gray как в PDF)
    hdr_texts = [
        "ID", "Наименование", "Изображение", "Кол-во,\nшт",
        "Цена за\nшт., руб.", "Сумма,\nруб.", "Магазин,\nссылка", "Примечание",
    ]
    for idx, cell in enumerate(table.rows[0].cells):
        _cell_bg(cell, 'D9D9D9')
        _cell_text(cell, hdr_texts[idx], bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    total_sum = 0
    IMG_W_CM = 2.5  # ширина картинки в ячейке

    for row_idx, item in enumerate(items):
        product = item.product
        quantity = item.quantity
        price = float(product.price)
        item_total = price * quantity
        total_sum += item_total

        row_cells = table.add_row().cells
        for idx, cell in enumerate(row_cells):
            cell.width = Cm(col_cms[idx])

        # Чередование фона строк (белый / светло-серый), как в PDF
        row_bg = 'FFFFFF' if row_idx % 2 == 0 else 'F7F7F7'
        for cell in row_cells:
            _cell_bg(cell, row_bg)

        # ID
        model_id = (
            product.model_3d_asset_ids.strip().split(',')[0]
            if getattr(product, 'model_3d_asset_ids', None) and product.model_3d_asset_ids.strip()
            else f'#{product.id}'
        )
        _cell_text(row_cells[0], model_id, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Наименование (без бренда)
        display_title = _strip_brand(product.title, getattr(product, 'brand', None))
        _cell_text(row_cells[1], display_title, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Изображение — конвертируем в JPEG для гарантированной совместимости с DOCX
        img_par = row_cells[2].paragraphs[0]
        img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_data = get_product_image(product)
        added_img = False
        if img_data:
            try:
                img_data.seek(0)
                pil_img = PILImage.open(img_data)
                pil_img.load()
                if pil_img.mode not in ('RGB', 'L'):
                    pil_img = pil_img.convert('RGB')
                jpeg_buf = io.BytesIO()
                pil_img.save(jpeg_buf, format='JPEG', quality=85)
                pil_img.close()
                jpeg_buf.seek(0)
                img_par.add_run().add_picture(jpeg_buf, width=Cm(IMG_W_CM))
                added_img = True
            except Exception:
                pass
        if not added_img:
            img_par.add_run('—')

        # Кол-во / Цена / Сумма
        _cell_text(row_cells[3], str(quantity), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row_cells[4], f'{price:,.0f}'.replace(',', ' '), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row_cells[5], f'{item_total:,.0f}'.replace(',', ' '), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Ссылка — поиск по полному оригинальному названию (с брендом = точные результаты)
        original_title = product.title or display_title
        search_query = quote_plus(f'{original_title} купить')
        search_url = f'https://ya.ru/search/?text={search_query}'
        link_label = f'ya.ru: {display_title[:30]}...' if len(display_title) > 30 else f'ya.ru: {display_title}'
        _add_hyperlink_to_cell(row_cells[6], search_url, link_label, docx_font, font_size=8)

        # Примечание (габариты + производитель)
        notes_parts = []
        dims = format_dimensions(product)
        if dims:
            notes_parts.append(dims.replace('\n', ' '))
        if getattr(product, 'cp_notes', None):
            notes_parts.append(product.cp_notes)
        elif getattr(product, 'brand', None):
            notes_parts.append(f'Производитель: {product.brand}')
        _cell_text(row_cells[7], ' '.join(notes_parts) if notes_parts else '—', size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- ИТОГО ----
    doc.add_paragraph()
    total_par = doc.add_paragraph()
    total_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(
        total_par.add_run(f'ИТОГО: {total_sum:,.0f} руб.'.replace(',', ' ')),
        docx_font, 12, bold=True,
    )

    # ---- ПРИМЕЧАНИЯ ----
    doc.add_paragraph()
    notes_p = doc.add_paragraph()
    _set_run_font(notes_p.add_run('Примечания:'), docx_font, 9, bold=True)
    _par('1. Смотреть совместно с планом расстановки мебели и развертками.')
    _par('2. Детальные чертежи для мебели индивидуального производства составлять совместно с поставщиками.')
    doc.add_paragraph()

    # ---- КАРТОЧКА ПРОЕКТА (справа внизу, как в PDF) ----
    card_wrap = doc.add_table(rows=1, cols=2)
    card_wrap.style = 'Table Grid'
    _remove_borders(card_wrap.rows[0].cells[0])
    card_wrap.rows[0].cells[0].width = Cm(21.2)
    right_cell = card_wrap.rows[0].cells[1]
    right_cell.width = Cm(5.5)

    inner = right_cell.add_table(rows=2, cols=1)
    inner.style = 'Table Grid'
    _cell_text(inner.rows[0].cells[0], 'КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ',
               bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(inner.rows[1].cells[0], f'Дизайн-проект: {proposal_request.project_name}',
               size=7, align=WD_ALIGN_PARAGRAPH.CENTER)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def send_proposal_email(proposal_request, pdf_bytes):
    """Отправляет КП по email"""
    from django.core.mail import EmailMessage
    
    subject = f"Коммерческое предложение - {proposal_request.project_name}"
    body = (
        f"Здравствуйте, {proposal_request.client_name}!\n\n"
        f"Направляем Вам коммерческое предложение по проекту «{proposal_request.project_name}».\n\n"
        f"Дата: {proposal_request.created_at.strftime('%d.%m.%Y')}\n"
    )
    if proposal_request.company_name:
        body += f"От: {proposal_request.company_name}\n"
    body += "\nС уважением,\nКоманда VIZHUB.PRO"
    
    email = EmailMessage(
        subject=subject,
        body=body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=[proposal_request.email],
    )
    
    filename = f"КП_{proposal_request.project_name}_{proposal_request.created_at.strftime('%Y%m%d')}.pdf"
    email.attach(filename, pdf_bytes, 'application/pdf')
    email.send()


def send_proposal_telegram(proposal_request, pdf_bytes):
    """Отправляет КП в Telegram"""
    import requests
    
    bot_token = getattr(settings, 'TELEGRAM_BOT_TOKEN', '')
    if not bot_token:
        raise ValueError("TELEGRAM_BOT_TOKEN не настроен в settings.py")
    
    telegram_id = proposal_request.telegram.strip()
    
    # Убираем @ если есть
    if telegram_id.startswith('@'):
        telegram_id = telegram_id[1:]
    
    # Отправляем сообщение
    caption = (
        f"📋 Коммерческое предложение\n\n"
        f"Проект: {proposal_request.project_name}\n"
        f"Клиент: {proposal_request.client_name}\n"
        f"Дата: {proposal_request.created_at.strftime('%d.%m.%Y')}\n"
    )
    if proposal_request.company_name:
        caption += f"От: {proposal_request.company_name}\n"
    
    filename = f"КП_{proposal_request.project_name}_{proposal_request.created_at.strftime('%Y%m%d')}.pdf"
    
    url = f"https://api.telegram.org/bot{bot_token}/sendDocument"
    
    files = {
        'document': (filename, pdf_bytes, 'application/pdf'),
    }
    data = {
        'chat_id': telegram_id,
        'caption': caption,
    }
    
    response = requests.post(url, data=data, files=files, timeout=30)
    
    if response.status_code != 200:
        error_data = response.json() if response.headers.get('content-type', '').startswith('application/json') else {}
        raise ValueError(
            f"Ошибка отправки в Telegram: {error_data.get('description', response.text)}"
        )
